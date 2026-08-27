"""Generate the interview prep document for completed stages.

Regenerate after each stage:  python scripts/build_interview_doc.py

Model answers are grounded in this project's real measured numbers. That is
the point: an answer with a number in it is credible, and an answer without
one sounds like a tutorial.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = (Path(__file__).resolve().parent.parent
       / "docs" / "interview" / "Interview-Prep-Stage-0-1.docx")

INK = RGBColor(0x16, 0x18, 0x2A)
BLUE = RGBColor(0x34, 0x41, 0x8C)
ROSE = RGBColor(0x9B, 0x33, 0x55)
GREY = RGBColor(0x46, 0x4A, 0x63)

SQL_ANSWER = (
    "SELECT skill_id, COUNT(*) AS n FROM posting_skills ps "
    "JOIN postings p ON p.id = ps.posting_id "
    "WHERE p.role_id = 'sde1-backend' "
    "GROUP BY skill_id ORDER BY n DESC LIMIT 10; "
    "Then divide by the total posting count to turn it into a percentage."
)

SECTIONS = [
    ("Opening - they always start here", [
        (
            "Tell me about this project.",
            "The first 60 seconds decide how the rest of the interview goes. "
            "They want to know if you can explain your own work without rambling.",
            "It reads thousands of real job postings for a target role, works out "
            "which skills that role actually demands and how often, then compares "
            "that against your resume and tells you which skill to learn next, "
            "ranked by how many more jobs it would qualify you for. I built it "
            "because I needed it myself for placement prep.",
            "Starting with the tech stack. Listing Python, FastAPI, React and NLP "
            "tells them nothing about what it does.",
        ),
        (
            "Why build this instead of a resume-to-job matcher?",
            "Checking whether you chose the problem or copied a tutorial.",
            "A resume matcher compares you against one job description, which only "
            "tells you about that one job. I wanted to know what the market as a "
            "whole demands, so I aggregate across thousands of postings for a role. "
            "That changes the question from 'am I a fit for this job' to 'what "
            "should I learn next'.",
            "Because it seemed interesting. No problem statement.",
        ),
        (
            "What stage is it at right now?",
            "Testing honesty. Overclaiming is the fastest way to lose credibility.",
            "The data pipeline is complete: 7,593 entry-level postings loaded and "
            "cleaned. Skill extraction is next, then an accuracy evaluation before "
            "I build any interface. I front-loaded the data and ML work because "
            "that is where the project could actually fail.",
            "Claiming it is finished when it is not. They will ask for a demo link.",
        ),
    ]),

    ("Your data - expect the most questions here", [
        (
            "Where did your data come from?",
            "Data provenance. Many student projects use invented data or "
            "illegally scraped data.",
            "A public Kaggle dataset of Naukri.com postings, about 96,000 unique "
            "listings. I used a frozen snapshot committed alongside the project "
            "rather than a live scraper, so the demo cannot break mid-interview "
            "and I am not violating any site's terms of service.",
            "I scraped LinkedIn. That is against their terms and interviewers "
            "know it.",
        ),
        (
            "You considered LinkedIn too. Why did you not use it?",
            "This is your single best story. It shows you measure instead of "
            "assume, which is rare in a fresher.",
            "I chose LinkedIn first because its descriptions are about ten times "
            "richer: median 3,435 characters against Naukri's 330, and 60 percent "
            "separate required from preferred skills. But when I counted postings "
            "per role, it had 4 entry-level backend jobs against Naukri's 259. It "
            "was a general job board dump, mostly retail and admin. So I switched. "
            "Naukri also matches the Indian market I am actually applying into.",
            "Any answer without numbers. The numbers are what make this credible.",
        ),
        (
            "How many postings do you have, and is that enough?",
            "Whether you thought about statistical validity at all.",
            "7,593 after cleaning, across three roles: 5,518 backend, 1,654 "
            "frontend, 421 data analyst. I set a minimum of 250 per role before I "
            "would compute percentages, and enforced it in code so the loader exits "
            "with an error if a role falls short. Below that, percentages look "
            "confident but are noise.",
            "Quoting the raw row count without mentioning cleaning, or having no "
            "threshold at all.",
        ),
        (
            "Walk me through how you cleaned the data.",
            "Very common. Real data work is mostly cleaning and they know it.",
            "Six steps, and the row count drops at each one: 103,657 raw, down to "
            "96,659 after removing repeated job IDs, 37,025 after filtering to 0 to "
            "2 years experience, 8,046 after matching titles to my three roles, "
            "7,860 after dropping descriptions under 60 characters, and 7,593 after "
            "content-based deduplication.",
            "I used pandas to clean it. No specifics.",
        ),
        (
            "What is deduplication and why did it matter here?",
            "Probing whether you understand the consequence, not just the term.",
            "Job boards repost the same listing under new IDs. If I count all of "
            "them, a statistic like '68 percent of jobs want Docker' might really "
            "mean one company posted the same ad 400 times. The dangerous part is "
            "that the number still looks completely believable. So I hash the "
            "title, company and first 500 characters of the description and drop "
            "repeats, within a source rather than across sources, because the same "
            "role appearing on two boards is genuinely two data points.",
            "Defining the word without explaining why it corrupts your results.",
        ),
        (
            "The dataset already had a skills column. Why extract anything?",
            "This question sinks people who used a pre-labelled dataset without "
            "inspecting it.",
            "I checked, and roughly half those tags are not skills. For backend "
            "roles the top tags include 'backend', 'coding', 'software', "
            "'development' and 'computer science': a job title, a meaningless verb "
            "and a degree. If I shipped that, my tool would advise people to learn "
            "'development'. So I built a skill taxonomy to normalise them, and kept "
            "the original tags in a separate table so I can measure my extraction "
            "against theirs.",
            "The dataset had them so I used them. You have measured evidence here, "
            "so use it.",
        ),
        (
            "Why did you filter to 0 to 2 years of experience?",
            "Whether your scope is deliberate.",
            "The tool is for freshers, so senior postings would skew the demand "
            "profile towards skills you cannot realistically have yet. Naukri gives "
            "a range like '0-5 Yrs' and I keep anything starting at 0, 1 or 2. I "
            "note in my limitations that a '0-5 Yrs' posting is a looser fit than a "
            "pure graduate role.",
            "Having no reason, or not knowing the limitation of your own filter.",
        ),
        (
            "Why exclude PHP and WordPress roles from backend?",
            "Tests whether you understand your own domain.",
            "The Indian entry-level market has a large PHP, WordPress and agency "
            "segment whose skill profile is genuinely different from "
            "product-company backend work. Pooling them produces a profile that "
            "describes neither accurately. They are a legitimate category, just a "
            "separate one.",
            "Not knowing the exclusions exist, or calling those roles inferior.",
        ),
    ]),

    ("Architecture and code structure", [
        (
            "Explain how your project is organised.",
            "Standard. They want to see whether you separated concerns or wrote "
            "one large file.",
            "Five parts. analyzer/ is the core logic, api/ is a thin FastAPI "
            "wrapper, web/ will be a single-page React frontend, data/ holds the "
            "taxonomy and frozen dataset, and eval/ holds my hand-labelled test set "
            "and scoring script.",
            "I have a main.py and some other files.",
        ),
        (
            "Why does your analyzer package not import FastAPI?",
            "If they notice this, they are impressed. Have the answer ready.",
            "Deliberate. The core has no web framework imports at all, so I could "
            "build and test the whole pipeline by running a Python file with no "
            "server involved. It also means the logic cannot get tangled up in "
            "request handling, and the API ended up as a thin wrapper with no "
            "business logic in it.",
            "Not knowing it was a decision.",
        ),
        (
            "Which database did you choose and why?",
            "Testing whether you can justify a choice rather than copy one.",
            "Postgres for production, SQLite while building. The data is deeply "
            "relational, with postings, skills, aliases and join tables, so a "
            "document database like MongoDB would have fought me. SQLite is a "
            "single file with zero setup, and because I access it through "
            "SQLAlchemy, moving to Postgres is a one-line change to the connection "
            "URL.",
            "I used MongoDB because it is easy, with no thought about data shape.",
        ),
        (
            "What is a virtual environment and why use one?",
            "A basic check. Easy marks, and easy to fumble.",
            "An isolated folder holding this project's library versions so they "
            "cannot clash with other projects on the same machine. It also makes "
            "requirements.txt meaningful, because anyone can recreate my exact "
            "environment from it.",
            "Not knowing, or confusing it with a virtual machine.",
        ),
        (
            "How do you manage dependencies?",
            "Follow-up to the above.",
            "requirements.txt with lower bounds rather than exact pins. I "
            "originally pinned exact versions, but they predated Python 3.14 and "
            "had no prebuilt wheels, so pip tried to compile pandas from source and "
            "failed on a missing C++ toolchain. I also split the heavy ML "
            "dependency into a separate file because it pulls 2.5GB of PyTorch and "
            "is not needed yet.",
            "Not having a requirements file at all.",
        ),
    ]),

    ("Fundamentals they slip in alongside", [
        (
            "Write a SQL query for the ten most-requested skills for backend roles.",
            "Very likely. They want to see you can actually write SQL, not just "
            "talk about it.",
            SQL_ANSWER,
            "Forgetting GROUP BY, or not knowing how to write the join.",
        ),
        (
            "What is the difference between precision and recall?",
            "Guaranteed if you mention ML at all.",
            "Precision is: of the skills my program found, how many were actually "
            "correct. Recall is: of the skills genuinely present, how many did it "
            "find. In my case recall matters more, because missing a skill means "
            "failing to tell a user about a real gap, whereas a false positive is "
            "visible on screen and easy to dismiss.",
            "Mixing them up, or reciting the formula without saying which one "
            "matters for your project.",
        ),
        (
            "What makes a good commit message?",
            "Low-stakes but revealing, and they do scroll your commit list.",
            "It explains why a change was made, not what changed, because the diff "
            "already shows what. Mine say things like 'switch primary source to "
            "Naukri' with the reasoning underneath, so anyone reading the history "
            "can follow the decisions.",
            "Update, fix, final.",
        ),
    ]),

    ("Curveballs - the ones that catch people out", [
        (
            "Why not just use ChatGPT to extract the skills?",
            "Increasingly common, and a trap. They are testing whether you can "
            "reason about trade-offs rather than dismiss or worship the tool.",
            "For 7,500 postings an LLM would be slow, cost money per call, and I "
            "could not explain why it made any given decision. The taxonomy "
            "approach is deterministic, instant, free and auditable. That said, I "
            "plan to run my 40 labelled postings through an LLM as a comparison "
            "baseline. If it clearly beats my extractor, that is a genuine finding "
            "worth reporting rather than hiding.",
            "LLMs are bad, which is dismissive and wrong. Or: I did not think of it.",
        ),
        (
            "Your data is from December 2024. Is it still relevant?",
            "Testing intellectual honesty about limitations.",
            "Partially. Core skills like SQL, Git and Docker move slowly, so those "
            "percentages hold. Anything fast-moving would drift. It is a genuine "
            "limitation and I state it in the README. The honest framing is that "
            "the pipeline is source-agnostic, so refreshing the snapshot is a data "
            "task, not a code change.",
            "Claiming it does not matter.",
        ),
        (
            "How do you know your percentages are actually correct?",
            "The sharpest question on this list. Have a real answer.",
            "I do not know it fully yet, which is exactly why the next stage is an "
            "evaluation before I build anything else. I hand-label 40 postings, "
            "measure precision and recall against them, and I have set a gate of 70 "
            "percent recall before proceeding. If it fails, I fix extraction rather "
            "than build a nice interface on top of bad numbers.",
            "It works. You have no evidence yet, and saying so is the stronger "
            "answer.",
        ),
        (
            "What was the hardest part so far?",
            "They want a real answer, not a humble-brag.",
            "Discovering my chosen dataset was wrong after I had already planned "
            "around it. LinkedIn had far better text and I wanted it to work. "
            "Counting the postings and accepting that 4 was not enough meant "
            "redoing the plan, but finding that in week one rather than week four "
            "is precisely why I checked.",
            "Nothing was hard. Or a fake weakness like caring too much about "
            "quality.",
        ),
        (
            "What would you do differently?",
            "Self-awareness check.",
            "I would count postings per role before choosing a dataset, not after. "
            "I evaluated text quality first because that was the interesting "
            "question, when volume was the one that could actually kill the "
            "project.",
            "Nothing.",
        ),
    ]),
]

CLOSING_QUESTIONS = [
    "What does the first six months look like for someone joining this team?",
    "How is code reviewed here, and what does that process look like day to day?",
    "What is the balance between building new features and maintaining existing systems?",
    "What would make you say a fresher had done well in their first year?",
]

FACTS = [
    ("Raw rows read", "103,657"),
    ("Unique postings", "96,659"),
    ("Entry-level pool (0-2 yrs)", "37,025"),
    ("Loaded after cleaning", "7,593"),
    ("Backend / Frontend / Data Analyst", "5,518 / 1,654 / 421"),
    ("Employer skill tags stored", "58,093"),
    ("Minimum postings per role (enforced)", "250"),
    ("LinkedIn entry-level backend postings", "4"),
    ("Naukri entry-level backend postings", "259"),
    ("Median description, LinkedIn vs Naukri", "3,435 vs 330 chars"),
    ("Stage 3 recall gate", "0.70"),
]


def _check_shape() -> None:
    """Guard the thing ISC004 warns about: a missing comma between two prose
    strings would silently merge two fields and shift everything after it.
    Checking the shape at runtime is a stronger guarantee than the lint rule,
    which is why that rule is scoped off for this file in pyproject.toml.
    """
    for heading, items in SECTIONS:
        for i, item in enumerate(items):
            if len(item) != 4:
                raise SystemExit(
                    f"{heading!r} item {i} has {len(item)} fields, expected 4 "
                    "(question, why, model, weak) -- likely a missing comma."
                )


def _label_run(para, label, text, label_color, label_pt, text_pt,
               italic=False, text_color=None):
    r = para.add_run(label)
    r.bold = True
    r.font.size = Pt(label_pt)
    r.font.color.rgb = label_color
    r2 = para.add_run(text)
    r2.font.size = Pt(text_pt)
    r2.italic = italic
    if text_color is not None:
        r2.font.color.rgb = text_color


def main() -> None:
    _check_shape()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_heading("Interview Preparation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    r = p.add_run("Skill-Gap Analyzer and Career Path Recommender  |  Stages 0-1")
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE
    r.bold = True

    p = doc.add_paragraph()
    r = p.add_run(
        "Questions a recruiter or hiring engineer is likely to ask about the work "
        "completed so far, with model answers grounded in this project's real "
        "measured numbers. Do not memorise the wording. Know the reasoning, and "
        "know the numbers."
    )
    r.italic = True
    r.font.color.rgb = GREY

    for heading, items in SECTIONS:
        doc.add_heading(heading, level=1)
        for question, why, model, weak in items:
            pq = doc.add_paragraph()
            rq = pq.add_run(question)
            rq.bold = True
            rq.font.size = Pt(11)
            rq.font.color.rgb = INK

            pw = doc.add_paragraph()
            _label_run(pw, "Why they ask:  ", why, BLUE, 9, 9, text_color=GREY)

            pm = doc.add_paragraph()
            pm.paragraph_format.left_indent = Pt(18)
            _label_run(pm, "Model answer:  ", model, BLUE, 9.5, 10, italic=True)

            pr = doc.add_paragraph()
            pr.paragraph_format.left_indent = Pt(18)
            _label_run(pr, "Weak answer:  ", weak, ROSE, 9, 9, text_color=GREY)

    doc.add_heading("Questions you should ask them", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Ask at least two. Asking none reads as disinterest, and these are the "
        "cheapest marks in the whole interview."
    )
    r.italic = True
    r.font.color.rgb = GREY
    for q in CLOSING_QUESTIONS:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Numbers to have at your fingertips", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Fact"
    hdr[1].text = "Value"
    for key, value in FACTS:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"written: {OUT}")


if __name__ == "__main__":
    main()
